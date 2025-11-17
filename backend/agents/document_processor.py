"""
Unified Document Processor

Handles multiple document formats:
- PDF (.pdf)
- Word documents (.docx, .doc)
- Markdown (.md)
- Plain text (.txt)
- Excel (.xlsx, .xls)

All blocking I/O operations are wrapped in asyncio.to_thread() to prevent
greenlet_spawn errors with SQLAlchemy's async session.
"""
import asyncio
import io
import mimetypes
from pathlib import Path
from typing import Optional, Tuple
import fitz  # PyMuPDF
import chardet


class DocumentProcessor:
    """
    Process various document formats and extract text content

    Supports: PDF, Word, Markdown, Text, Excel
    """

    MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB

    # Supported file extensions and their MIME types
    SUPPORTED_FORMATS = {
        '.pdf': 'application/pdf',
        '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        '.doc': 'application/msword',
        '.txt': 'text/plain',
        '.md': 'text/markdown',
        '.markdown': 'text/markdown',
        '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        '.xls': 'application/vnd.ms-excel',
    }

    @staticmethod
    def _extract_pdf_sync(file_bytes: bytes) -> str:
        """Extract text from PDF (synchronous)"""
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text_parts = []

        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text()
            if text.strip():
                text_parts.append(f"--- Page {page_num + 1} ---\n{text}")

        doc.close()
        return "\n\n".join(text_parts)

    @staticmethod
    def _extract_docx_sync(file_bytes: bytes) -> str:
        """Extract text from Word document (synchronous)"""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx is required for Word document processing")

        doc = Document(io.BytesIO(file_bytes))
        text_parts = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)

        return "\n\n".join(text_parts)

    @staticmethod
    def _extract_excel_sync(file_bytes: bytes, filename: str) -> str:
        """Extract text from Excel file (synchronous)"""
        try:
            from openpyxl import load_workbook
        except ImportError:
            raise ImportError("openpyxl is required for Excel file processing")

        wb = load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
        text_parts = []

        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            text_parts.append(f"--- Sheet: {sheet_name} ---")

            for row in sheet.iter_rows(values_only=True):
                row_text = " | ".join(str(cell) if cell is not None else "" for cell in row)
                if row_text.strip():
                    text_parts.append(row_text)

        wb.close()
        return "\n\n".join(text_parts)

    @staticmethod
    def _extract_text_sync(file_bytes: bytes) -> str:
        """Extract text from plain text file with encoding detection (synchronous)"""
        # Detect encoding
        detected = chardet.detect(file_bytes)
        encoding = detected.get('encoding', 'utf-8')

        try:
            text = file_bytes.decode(encoding)
        except (UnicodeDecodeError, AttributeError):
            # Fallback to utf-8 with error handling
            text = file_bytes.decode('utf-8', errors='ignore')

        return text

    @staticmethod
    def get_file_extension(filename: str) -> str:
        """Get normalized file extension"""
        return Path(filename).suffix.lower()

    @staticmethod
    def is_supported(filename: str) -> bool:
        """Check if file format is supported"""
        ext = DocumentProcessor.get_file_extension(filename)
        return ext in DocumentProcessor.SUPPORTED_FORMATS

    @staticmethod
    def get_mime_type(filename: str) -> Optional[str]:
        """Get MIME type for file"""
        ext = DocumentProcessor.get_file_extension(filename)
        return DocumentProcessor.SUPPORTED_FORMATS.get(ext)

    @staticmethod
    async def validate_document(file_bytes: bytes, filename: str) -> Tuple[bool, str]:
        """
        Validate document file

        Args:
            file_bytes: File content
            filename: Original filename

        Returns:
            Tuple of (is_valid, error_message)
        """
        if not file_bytes or len(file_bytes) == 0:
            return False, "File is empty"

        if len(file_bytes) > DocumentProcessor.MAX_FILE_SIZE:
            max_mb = DocumentProcessor.MAX_FILE_SIZE / (1024 * 1024)
            return False, f"File too large. Maximum size: {max_mb:.1f}MB"

        if not DocumentProcessor.is_supported(filename):
            ext = DocumentProcessor.get_file_extension(filename)
            supported = ", ".join(DocumentProcessor.SUPPORTED_FORMATS.keys())
            return False, f"Unsupported file type: {ext}. Supported: {supported}"

        return True, ""

    @staticmethod
    async def extract_text(file_bytes: bytes, filename: str) -> str:
        """
        Extract text from document (async wrapper)

        Args:
            file_bytes: File content as bytes
            filename: Original filename

        Returns:
            Extracted text as string

        Raises:
            ValueError: If file is invalid
            Exception: If processing fails
        """
        # Validate
        is_valid, error_msg = await DocumentProcessor.validate_document(file_bytes, filename)
        if not is_valid:
            raise ValueError(error_msg)

        ext = DocumentProcessor.get_file_extension(filename)

        try:
            # Route to appropriate processor (run in thread pool to avoid blocking)
            if ext == '.pdf':
                text = await asyncio.to_thread(DocumentProcessor._extract_pdf_sync, file_bytes)
            elif ext in ['.docx', '.doc']:
                text = await asyncio.to_thread(DocumentProcessor._extract_docx_sync, file_bytes)
            elif ext in ['.xlsx', '.xls']:
                text = await asyncio.to_thread(DocumentProcessor._extract_excel_sync, file_bytes, filename)
            elif ext in ['.txt', '.md', '.markdown']:
                text = await asyncio.to_thread(DocumentProcessor._extract_text_sync, file_bytes)
            else:
                raise ValueError(f"Unsupported file extension: {ext}")

            if not text or not text.strip():
                raise ValueError("No text content found in document")

            return text

        except ValueError:
            raise
        except ImportError as e:
            raise Exception(f"Missing dependency: {str(e)}")
        except Exception as e:
            raise Exception(f"Document processing failed: {str(e)}")

    @staticmethod
    async def get_document_info(file_bytes: bytes, filename: str) -> dict:
        """
        Get document metadata

        Args:
            file_bytes: File content
            filename: Original filename

        Returns:
            Dictionary with document metadata
        """
        ext = DocumentProcessor.get_file_extension(filename)
        mime_type = DocumentProcessor.get_mime_type(filename)

        info = {
            "filename": filename,
            "extension": ext,
            "mime_type": mime_type,
            "size_bytes": len(file_bytes),
            "size_mb": len(file_bytes) / (1024 * 1024),
            "is_supported": DocumentProcessor.is_supported(filename)
        }

        # Try to get page count for PDFs
        if ext == '.pdf':
            try:
                def get_page_count():
                    doc = fitz.open(stream=file_bytes, filetype="pdf")
                    count = len(doc)
                    doc.close()
                    return count

                info["page_count"] = await asyncio.to_thread(get_page_count)
            except Exception:
                info["page_count"] = None

        return info
